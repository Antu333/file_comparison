import io
import csv

import docx
import pandas as pd
from pypdf import PdfReader
from fastapi import UploadFile


class FileReader:

    # ── DOCX ──────────────────────────────────────────────────────────────────

    @staticmethod
    async def get_text(upload: UploadFile) -> str:
        contents = await upload.read()
        doc = docx.Document(io.BytesIO(contents))
        return "\n".join(p.text for p in doc.paragraphs)

    @staticmethod
    def get_text_from_path(path: str) -> str:
        doc = docx.Document(path)
        return "\n".join(p.text for p in doc.paragraphs)

    # ── PDF ───────────────────────────────────────────────────────────────────

    @staticmethod
    async def extract_text_from_pdf(upload: UploadFile) -> str:
        contents = await upload.read()
        reader = PdfReader(io.BytesIO(contents))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    @staticmethod
    def extract_text_from_pdf_path(path: str) -> str:
        reader = PdfReader(path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

    # ── CSV ───────────────────────────────────────────────────────────────────

    @staticmethod
    async def read_csv(upload: UploadFile) -> pd.DataFrame:
        contents = await upload.read()
        return pd.read_csv(io.BytesIO(contents))

    @staticmethod
    def read_csv_from_path(path: str) -> pd.DataFrame:
        return pd.read_csv(path)

    # ── Excel ─────────────────────────────────────────────────────────────────

    @staticmethod
    async def read_excel(upload: UploadFile) -> pd.DataFrame:
        contents = await upload.read()
        return pd.read_excel(io.BytesIO(contents))

    @staticmethod
    def read_excel_from_path(path: str) -> pd.DataFrame:
        return pd.read_excel(path)

    # ── Helpers ───────────────────────────────────────────────────────────────

    @staticmethod
    def read_csv_string(csv_str: str) -> list[dict]:
        """Parse a raw CSV string (from Gemini) into a list of dicts."""
        s = io.StringIO(csv_str)
        reader = csv.DictReader(s)
        processed = []
        for row in reader:
            processed.append({
                "sl_no": row.get("sl_no", ""),
                "column_name": row.get("column_name", ""),
                "old_value": row.get("old_value", ""),
                "new_value": row.get("new_value", ""),
            })
        return processed

    @staticmethod
    async def read_to_tempfile(upload: UploadFile, suffix: str) -> bytes:
        """Read an UploadFile to raw bytes (for callers that need a real path)."""
        return await upload.read()